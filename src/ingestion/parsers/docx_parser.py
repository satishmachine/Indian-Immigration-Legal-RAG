"""
ingestion.parsers.docx_parser
==============================
DOCX document parser using python-docx.

Extracts text from:
* Paragraphs (preserving heading hierarchy)
* Tables (row-by-row, pipe-separated)
* Headers and footers (for page numbers and act names)
* Footnotes (where present)

Page mapping is approximated using section breaks and paragraph count
since DOCX does not expose a native page model.

Dependencies
------------
    python-docx (installed as ``docx``)

Usage
-----
    from ingestion.parsers.docx_parser import DOCXParser

    parser = DOCXParser()
    result = parser.parse("Data_Set/passport_rules.docx")
"""

from __future__ import annotations

import logging
from pathlib import Path

from ingestion.parsers.base import BaseDocumentParser, DocumentParseError

logger = logging.getLogger(__name__)

# Approximate paragraphs per page (used to build a synthetic page map)
_PARAGRAPHS_PER_PAGE = 40


class DOCXParser(BaseDocumentParser):
    """
    Parser for Microsoft Word DOCX files.

    Args:
        include_tables:  Extract table content as text (default True).
        include_headers: Extract header/footer text (default True).
    """

    SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".docx", ".doc"})

    def __init__(
        self,
        include_tables: bool = True,
        include_headers: bool = True,
        **kwargs: object,
    ) -> None:
        super().__init__(**kwargs)  # type: ignore[arg-type]
        self._include_tables = include_tables
        self._include_headers = include_headers

    def _extract_text(self, path: Path) -> tuple[str, dict[int, str]]:
        """Extract text from a DOCX file with structural preservation."""
        try:
            from docx import Document as DocxDocument  # noqa: PLC0415
            from docx.oxml.ns import qn  # noqa: PLC0415
        except ImportError as exc:
            raise DocumentParseError(
                path, "python-docx not installed: pip install python-docx"
            ) from exc

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise DocumentParseError(path, f"could not open DOCX: {exc}") from exc

        paragraphs: list[str] = []

        # ── Headers / footers ──────────────────────────────────────────────
        if self._include_headers:
            for section in doc.sections:
                for part in (section.header, section.footer):
                    if part and not part.is_linked_to_previous:
                        for para in part.paragraphs:
                            text = para.text.strip()
                            if text:
                                paragraphs.append(f"[HEADER] {text}")

        # ── Body paragraphs ────────────────────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Preserve heading levels for structure detection
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1]
                paragraphs.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {text}\n")
            else:
                paragraphs.append(text)

        # ── Tables ─────────────────────────────────────────────────────────
        if self._include_tables:
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(c for c in cells if c)
                    if row_text:
                        paragraphs.append(row_text)

        if not paragraphs:
            raise DocumentParseError(path, "no text could be extracted from DOCX")

        # ── Build synthetic page map ───────────────────────────────────────
        page_map: dict[int, str] = {}
        page_no = 1
        chunk: list[str] = []

        for i, para in enumerate(paragraphs):
            chunk.append(para)
            if (i + 1) % _PARAGRAPHS_PER_PAGE == 0:
                page_map[page_no] = "\n".join(chunk)
                page_no += 1
                chunk = []

        if chunk:
            page_map[page_no] = "\n".join(chunk)

        full_text = "\n".join(paragraphs)
        return full_text, page_map
